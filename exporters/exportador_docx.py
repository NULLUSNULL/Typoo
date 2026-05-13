# exporters/exportador_docx.py
# Exportación del manuscrito al formato Word (.docx) usando python-docx

from __future__ import annotations

import re
from pathlib import Path
from typing import Optional

from core.logger import logger
from models.proyecto import Proyecto


class ExportadorDocx:
    """
    Genera un archivo .docx con el contenido del proyecto literario,
    convirtiendo la estructura Markdown a estilos de Word.
    Requiere la dependencia: python-docx
    """

    @staticmethod
    def exportar(
        proyecto: Proyecto,
        ruta_destino: Path,
        incluir_portada: bool = True,
    ) -> bool:
        """Exporta el proyecto completo a un archivo .docx."""
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            logger.error("python-docx no está instalado. Ejecuta: pip install python-docx")
            return False

        try:
            doc = Document()
            ExportadorDocx._configurar_estilos(doc)

            if incluir_portada:
                ExportadorDocx._agregar_portada(doc, proyecto)

            ExportadorDocx._recorrer_arbol(doc, proyecto)

            doc.save(str(ruta_destino))
            logger.info("Exportado DOCX: %s", ruta_destino)
            return True
        except Exception as e:
            logger.error("Error al exportar DOCX: %s", e)
            return False

    @staticmethod
    def _configurar_estilos(doc) -> None:
        """Ajusta márgenes y estilos básicos del documento Word."""
        from docx.shared import Inches, Pt
        from docx.oxml.ns import qn

        for seccion in doc.sections:
            seccion.top_margin    = Inches(1.2)
            seccion.bottom_margin = Inches(1.2)
            seccion.left_margin   = Inches(1.4)
            seccion.right_margin  = Inches(1.4)

        estilo_normal = doc.styles["Normal"]
        estilo_normal.font.name = "Garamond"
        estilo_normal.font.size = Pt(12)
        estilo_normal.paragraph_format.line_spacing = Pt(22)

    @staticmethod
    def _agregar_portada(doc, proyecto: Proyecto) -> None:
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for _ in range(8):
            doc.add_paragraph()

        titulo = doc.add_heading(proyecto.nombre, level=0)
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if proyecto.autor:
            autor = doc.add_paragraph(f"por {proyecto.autor}")
            autor.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_page_break()

    @staticmethod
    def _recorrer_arbol(doc, proyecto: Proyecto) -> None:
        from models.documento import ItemProyecto

        def procesar(item: ItemProyecto) -> None:
            if item.ruta_relativa:
                ruta = proyecto.ruta / item.ruta_relativa
                if ruta.is_file():
                    texto = ruta.read_text(encoding="utf-8")
                    ExportadorDocx._markdown_a_docx(doc, texto)

            for hijo in sorted(item.hijos, key=lambda h: h.orden):
                procesar(hijo)

        if proyecto.raiz:
            for hijo in sorted(proyecto.raiz.hijos, key=lambda h: h.orden):
                procesar(hijo)

    @staticmethod
    def _markdown_a_docx(doc, texto: str) -> None:
        """Convierte líneas Markdown a párrafos y encabezados de Word."""
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        lineas = texto.splitlines()
        en_bloque_codigo = False
        buffer_codigo: list[str] = []

        for linea in lineas:
            # Bloque de código
            if linea.startswith("```"):
                if en_bloque_codigo:
                    doc.add_paragraph("\n".join(buffer_codigo), style="No Spacing")
                    buffer_codigo.clear()
                en_bloque_codigo = not en_bloque_codigo
                continue

            if en_bloque_codigo:
                buffer_codigo.append(linea)
                continue

            # Encabezados
            m_enc = re.match(r"^(#{1,6})\s+(.+)$", linea)
            if m_enc:
                nivel = len(m_enc.group(1))
                doc.add_heading(m_enc.group(2), level=nivel)
                continue

            # Separador horizontal
            if re.match(r"^[-*_]{3,}\s*$", linea):
                p = doc.add_paragraph("* * *")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                continue

            # Cita
            if linea.startswith("> "):
                p = doc.add_paragraph(linea[2:], style="Quote")
                continue

            # Lista
            if re.match(r"^\s*[-*+]\s", linea):
                doc.add_paragraph(re.sub(r"^\s*[-*+]\s", "", linea), style="List Bullet")
                continue

            if re.match(r"^\s*\d+\.\s", linea):
                doc.add_paragraph(re.sub(r"^\s*\d+\.\s", "", linea), style="List Number")
                continue

            # Párrafo normal con formato inline
            if linea.strip():
                p = doc.add_paragraph()
                ExportadorDocx._aplicar_formato_inline(p, linea)
            else:
                doc.add_paragraph()

    @staticmethod
    def _aplicar_formato_inline(parrafo, texto: str) -> None:
        """Aplica negrita e cursiva inline dentro de un párrafo Word."""
        # Patrón simplificado para negrita y cursiva
        patron = re.compile(r"(\*\*|__|\*|_|~~)")
        partes = patron.split(texto)
        en_negrita = False
        en_cursiva = False
        en_tachado = False

        for parte in partes:
            if parte == "**" or parte == "__":
                en_negrita = not en_negrita
            elif parte == "*" or parte == "_":
                en_cursiva = not en_cursiva
            elif parte == "~~":
                en_tachado = not en_tachado
            else:
                run = parrafo.add_run(parte)
                run.bold = en_negrita
                run.italic = en_cursiva
