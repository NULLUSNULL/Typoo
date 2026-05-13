"""
Export service for converting documents to various formats.
"""

import re
from pathlib import Path
from typing import Optional

try:
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
    from reportlab.lib.units import inch
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

from typoo.logger import Logger
from typoo.models import Project, ExportFormat
from typoo.utils import MarkdownUtils

logger = Logger.get()


class ExportService:
    """Service for exporting projects to various formats."""
    
    @staticmethod
    def export_project(project: Project, output_path: Path, format: ExportFormat) -> bool:
        """
        Export project to specified format.
        
        Args:
            project: Project to export
            output_path: Path for output file
            format: Export format
            
        Returns:
            True if successful, False otherwise
        """
        try:
            if format == ExportFormat.DOCX:
                return ExportService._export_to_docx(project, output_path)
            elif format == ExportFormat.PDF:
                return ExportService._export_to_pdf(project, output_path)
            elif format == ExportFormat.TXT:
                return ExportService._export_to_txt(project, output_path)
            elif format == ExportFormat.MARKDOWN:
                return ExportService._export_to_markdown(project, output_path)
            else:
                logger.warning(f"Unsupported export format: {format}")
                return False
        except Exception as e:
            logger.error(f"Export failed: {str(e)}")
            return False
    
    @staticmethod
    def _export_to_docx(project: Project, output_path: Path) -> bool:
        """Export project to DOCX format."""
        if not DOCX_AVAILABLE:
            logger.error("python-docx is not installed")
            return False
        
        try:
            doc = DocxDocument()
            
            # Add title
            title = doc.add_heading(project.metadata.name, level=0)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # Add author if available
            if project.metadata.author:
                author = doc.add_paragraph(f"Author: {project.metadata.author}")
                author.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # Add documents
            for document in project.documents:
                ExportService._add_document_to_docx(doc, document)
            
            doc.save(output_path)
            logger.info(f"Project exported to DOCX: {output_path}")
            return True
        except Exception as e:
            logger.error(f"DOCX export error: {str(e)}")
            return False
    
    @staticmethod
    def _add_document_to_docx(doc: DocxDocument, document) -> None:
        """Add document to DOCX recursively."""
        from typoo.models import DocumentType
        
        # Add heading based on document type
        level = 1 if document.doc_type == DocumentType.CHAPTER else 2
        heading = doc.add_heading(document.name, level=level)
        
        # Add content
        if document.content:
            paragraphs = document.content.split('\n')
            for para in paragraphs:
                if para.strip():
                    doc.add_paragraph(para)
        
        # Add children
        for child in document.children:
            ExportService._add_document_to_docx(doc, child)
    
    @staticmethod
    def _export_to_pdf(project: Project, output_path: Path) -> bool:
        """Export project to PDF format."""
        if not PDF_AVAILABLE:
            logger.error("reportlab is not installed")
            return False
        
        try:
            story = []
            styles = getSampleStyleSheet()
            
            # Add title
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                textColor=RGBColor(0, 0, 0),
                spaceAfter=30,
                alignment=1  # CENTER
            )
            story.append(Paragraph(project.metadata.name, title_style))
            
            # Add content from all documents
            for document in project.documents:
                ExportService._add_document_to_pdf(story, document, styles)
            
            # Build PDF
            doc = SimpleDocTemplate(str(output_path), pagesize=letter)
            doc.build(story)
            logger.info(f"Project exported to PDF: {output_path}")
            return True
        except Exception as e:
            logger.error(f"PDF export error: {str(e)}")
            return False
    
    @staticmethod
    def _add_document_to_pdf(story: list, document, styles) -> None:
        """Add document to PDF story recursively."""
        from typoo.models import DocumentType
        
        # Add heading
        level = 'Heading1' if document.doc_type == DocumentType.CHAPTER else 'Heading2'
        story.append(Paragraph(document.name, styles[level]))
        story.append(Spacer(1, 0.2 * inch))
        
        # Add content
        if document.content:
            story.append(Paragraph(document.content, styles['Normal']))
            story.append(Spacer(1, 0.3 * inch))
        
        # Add children
        for child in document.children:
            ExportService._add_document_to_pdf(story, child, styles)
    
    @staticmethod
    def _export_to_txt(project: Project, output_path: Path) -> bool:
        """Export project to plain text format."""
        try:
            content = []
            content.append(f"{project.metadata.name}\n")
            if project.metadata.author:
                content.append(f"Author: {project.metadata.author}\n")
            content.append("=" * 80 + "\n\n")
            
            # Collect all document content
            for document in project.documents:
                ExportService._collect_text_content(content, document, level=0)
            
            with open(output_path, "w", encoding="utf-8") as f:
                f.writelines(content)
            
            logger.info(f"Project exported to TXT: {output_path}")
            return True
        except Exception as e:
            logger.error(f"TXT export error: {str(e)}")
            return False
    
    @staticmethod
    def _collect_text_content(content: list, document, level: int = 0) -> None:
        """Recursively collect text content from documents."""
        indent = "  " * level
        content.append(f"{indent}{document.name}\n")
        if level == 0:
            content.append("-" * 40 + "\n")
        
        if document.content:
            content.append(document.content + "\n\n")
        
        for child in document.children:
            ExportService._collect_text_content(content, child, level + 1)
    
    @staticmethod
    def _export_to_markdown(project: Project, output_path: Path) -> bool:
        """Export project to Markdown format."""
        try:
            content = []
            content.append(f"# {project.metadata.name}\n\n")
            if project.metadata.author:
                content.append(f"**Author:** {project.metadata.author}\n\n")
            if project.metadata.description:
                content.append(f"{project.metadata.description}\n\n")
            content.append("---\n\n")
            
            # Collect markdown content
            for document in project.documents:
                ExportService._collect_markdown_content(content, document, level=1)
            
            with open(output_path, "w", encoding="utf-8") as f:
                f.writelines(content)
            
            logger.info(f"Project exported to Markdown: {output_path}")
            return True
        except Exception as e:
            logger.error(f"Markdown export error: {str(e)}")
            return False
    
    @staticmethod
    def _collect_markdown_content(content: list, document, level: int = 1) -> None:
        """Recursively collect markdown content from documents."""
        heading = "#" * min(level, 6)
        content.append(f"{heading} {document.name}\n\n")
        
        if document.content:
            content.append(document.content + "\n\n")
        
        for child in document.children:
            ExportService._collect_markdown_content(content, child, level + 1)
